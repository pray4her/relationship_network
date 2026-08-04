"""Company document type validation and text extraction."""

from __future__ import annotations

import hashlib
import io
import zipfile
from dataclasses import dataclass
from pathlib import PurePosixPath
from typing import Final, Literal, final

MAX_DOCUMENT_BYTES: Final = 10 * 1024 * 1024
ALLOWED_EXTENSIONS: Final[frozenset[str]] = frozenset({".pdf", ".docx", ".txt"})
DocumentKind = Literal["pdf", "docx", "txt"]

INVALID_DOCUMENT_DETAIL: Final = "invalid_document"
DOCUMENT_TOO_LARGE_DETAIL: Final = "document_too_large"


@final
class InvalidDocumentError(Exception):
    """Raised when an upload fails type or content validation."""


@final
class DocumentTooLargeError(Exception):
    """Raised when an upload exceeds the 10 MB limit."""


@final
@dataclass(frozen=True)
class ValidatedDocument:
    """A validated company document ready for storage and extraction."""

    kind: DocumentKind
    content_type: str
    original_filename: str
    data: bytes
    sha256: str
    scan_status: Literal["content_checked"] = "content_checked"


def validate_document(*, filename: str, data: bytes) -> ValidatedDocument:
    """Validate size, extension, and magic bytes; reject executables and spoofs."""
    if len(data) > MAX_DOCUMENT_BYTES:
        raise DocumentTooLargeError
    if len(data) == 0:
        raise InvalidDocumentError
    if _looks_like_executable(data):
        raise InvalidDocumentError
    extension = PurePosixPath(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise InvalidDocumentError
    kind = _detect_kind(data, extension)
    content_type = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain; charset=utf-8",
    }[kind]
    return ValidatedDocument(
        kind=kind,
        content_type=content_type,
        original_filename=PurePosixPath(filename).name,
        data=data,
        sha256=hashlib.sha256(data).hexdigest(),
    )


def extract_text(document: ValidatedDocument) -> str:
    """Extract plain text from a validated company document."""
    if document.kind == "txt":
        return _decode_text(document.data)
    if document.kind == "pdf":
        return _extract_pdf(document.data)
    return _extract_docx(document.data)


def _detect_kind(data: bytes, extension: str) -> DocumentKind:
    if extension == ".pdf":
        if not data.lstrip().startswith(b"%PDF"):
            raise InvalidDocumentError
        return "pdf"
    if extension == ".docx":
        if not _is_docx(data):
            raise InvalidDocumentError
        return "docx"
    # .txt — require decodable text without NUL
    if b"\x00" in data:
        raise InvalidDocumentError
    try:
        _ = data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            _ = data.decode("gb18030")
        except UnicodeDecodeError as error:
            raise InvalidDocumentError from error
    return "txt"


def _is_docx(data: bytes) -> bool:
    if data[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            names = set(archive.namelist())
    except zipfile.BadZipFile:
        return False
    return "[Content_Types].xml" in names and any(
        name.startswith("word/") for name in names
    )


def _looks_like_executable(data: bytes) -> bool:
    return data.startswith((b"MZ", b"\x7fELF", b"\xcf\xfa\xed\xfe", b"\xfe\xed\xfa\xce"))


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("gb18030", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text).strip()
